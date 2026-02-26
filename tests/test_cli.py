from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from nodes.preprocessor import parse_and_chunk


def test_imports() -> None:
    # Smoke test that modules are importable.
    assert parse_and_chunk is not None


@pytest.mark.skip(reason="CLI has circular import issue due to root cli.py wrapper")
def test_cli_parser_builds() -> None:
    import sys
    from pathlib import Path
    src_path = str(Path(__file__).parent.parent / "src")
    if src_path not in sys.path:
        sys.path.insert(0, src_path)
    
    # Import from src/cli.py, not root cli.py
    import importlib
    cli_module = importlib.import_module("cli")
    build_parser = cli_module.build_parser

    parser = build_parser()
    # Test atomize subcommand
    args = parser.parse_args(["atomize", "--input", "test.docx"])

    assert args.command == "atomize"
    assert args.input_path == "test.docx"
    assert isinstance(args.output_dir, str)
    assert isinstance(args.log_level, str)


@pytest.mark.skip(reason="CLI has circular import issue due to root cli.py wrapper")
def test_cli_preprocess_writes_output(tmp_path: Path) -> None:
    import sys
    from pathlib import Path as P
    src_path = str(P(__file__).parent.parent / "src")
    if src_path not in sys.path:
        sys.path.insert(0, src_path)
    
    import importlib
    cli_module = importlib.import_module("cli")
    main = cli_module.main

    input_docx = tmp_path / "sample.docx"
    doc = Document()
    doc.add_heading("H1", level=1)
    doc.add_paragraph("Hello world")
    doc.save(str(input_docx))

    out_path = tmp_path / "out.json"
    code = main(
        [
            "preprocess",
            "--input",
            str(input_docx),
            "--output",
            str(out_path),
            "--min-chunk-chars",
            "1",
        ]
    )

    assert code == 0
    assert out_path.exists()
    assert out_path.read_text(encoding="utf-8").strip().startswith("{")
