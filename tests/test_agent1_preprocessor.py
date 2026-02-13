from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from agent1.exceptions import EmptyDocumentError
from agent1.nodes.preprocessor import parse_and_chunk


def _write_docx(
    path: Path,
    *,
    with_content: bool = True,
    large_table: bool = False,
    large_kv_table: bool = False,
    stub_control: bool = False,
    policy_and_risk: bool = False,
) -> None:
    doc = Document()

    if not with_content:
        doc.save(str(path))
        return

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph("This is a short executive summary paragraph.")

    doc.add_heading("Policy Library", level=1)

    p1 = doc.add_paragraph("Item 1")
    p1.style = "List Bullet"
    p2 = doc.add_paragraph("Item 2")
    p2.style = "List Bullet"

    doc.add_paragraph("This is a prose paragraph one.")
    doc.add_paragraph("This is a prose paragraph two.")

    # Small table
    t = doc.add_table(rows=3, cols=2)
    t.cell(0, 0).text = "Field"
    t.cell(0, 1).text = "Value"
    t.cell(1, 0).text = "Policy ID"
    t.cell(1, 1).text = "P-001"
    t.cell(2, 0).text = "Owner"
    t.cell(2, 1).text = "Compliance"

    if large_kv_table:
        # Extend the same vertical key-value table to exceed max_chunk_chars.
        for i in range(80):
            r = t.add_row().cells
            r[0].text = f"Field {i}"
            r[1].text = ("X" * 40) + str(i)

    if large_table:
        t2 = doc.add_table(rows=1, cols=3)
        t2.cell(0, 0).text = "ColA"
        t2.cell(0, 1).text = "ColB"
        t2.cell(0, 2).text = "ColC"
        for i in range(120):
            r = t2.add_row().cells
            r[0].text = f"A{i}"
            r[1].text = ("B" * 30) + str(i)
            r[2].text = ("C" * 30) + str(i)

    if stub_control:
        doc.add_heading("C-011: Stub Control", level=2)
        doc.add_paragraph("Full control details documented in RSA Archer Controls module.")

    if policy_and_risk:
        doc.add_heading("P-001: Test Policy", level=2)
        t_p = doc.add_table(rows=2, cols=2)
        t_p.cell(0, 0).text = "Policy ID"
        t_p.cell(0, 1).text = "P-001"
        t_p.cell(1, 0).text = "Description"
        t_p.cell(1, 1).text = "Test policy description"

        doc.add_heading("R-001: Test Risk", level=2)
        t_r = doc.add_table(rows=2, cols=2)
        t_r.cell(0, 0).text = "Risk ID"
        t_r.cell(0, 1).text = "R-001"
        t_r.cell(1, 0).text = "Description"
        t_r.cell(1, 1).text = "Test risk description"

    doc.save(str(path))


def test_docx_heading_tracking_and_types(tmp_path: Path) -> None:
    path = tmp_path / "sample.docx"
    _write_docx(path, with_content=True)

    out = parse_and_chunk(path, file_type="docx", max_chunk_chars=500, min_chunk_chars=1)

    assert out.file_type == "docx"
    assert out.total_chunks == len(out.chunks)
    assert out.document_stats["heading_count"] >= 2
    assert out.document_stats["table_count"] >= 1
    assert not any(c.chunk_type == "heading" for c in out.chunks)
    assert any(c.chunk_type == "list" for c in out.chunks)
    assert any(c.chunk_type == "prose" for c in out.chunks)
    assert any(c.chunk_type == "table" for c in out.chunks)

    # Headings are not emitted as chunks; they must be attached via parent_heading.
    assert any(c.parent_heading == "Policy Library" for c in out.chunks)


def test_docx_table_splitting_repeats_header(tmp_path: Path) -> None:
    path = tmp_path / "large_table.docx"
    _write_docx(path, with_content=True, large_table=True)

    out = parse_and_chunk(path, file_type="docx", max_chunk_chars=250, min_chunk_chars=1)

    table_chunks = [c for c in out.chunks if c.chunk_type == "table" and c.table_data]
    large_table_chunks = [c for c in table_chunks if c.col_count == 3]
    assert len(large_table_chunks) >= 2

    header = large_table_chunks[0].table_data[0]
    assert header == ["ColA", "ColB", "ColC"]

    for tc in large_table_chunks:
        assert tc.table_data[0] == header
        assert tc.row_count is not None and tc.row_count >= 1
        assert tc.col_count == 3


def test_docx_kv_table_never_splits_mid_entity(tmp_path: Path) -> None:
    path = tmp_path / "large_kv.docx"
    _write_docx(path, with_content=True, large_kv_table=True)

    out = parse_and_chunk(path, file_type="docx", max_chunk_chars=250, min_chunk_chars=1)

    # The 2-column vertical key-value table should remain a single chunk.
    kv_chunks = [c for c in out.chunks if c.chunk_type == "table" and c.col_count == 2]
    assert len(kv_chunks) == 1
    assert kv_chunks[0].row_count is not None and kv_chunks[0].row_count > 10


def test_docx_stub_control_flagged_incomplete(tmp_path: Path) -> None:
    path = tmp_path / "stub_control.docx"
    _write_docx(path, with_content=True, stub_control=True)

    out = parse_and_chunk(path, file_type="docx", max_chunk_chars=1000, min_chunk_chars=1)

    stub_chunks = [c for c in out.chunks if c.annotations.get("incomplete_record") is True]
    assert len(stub_chunks) >= 1
    assert any(c.annotations.get("record_id") == "C-011" for c in stub_chunks)


def test_chunk_id_determinism(tmp_path: Path) -> None:
    path = tmp_path / "determinism.docx"
    _write_docx(path, with_content=True)

    out1 = parse_and_chunk(path, file_type="docx", max_chunk_chars=500, min_chunk_chars=1)
    out2 = parse_and_chunk(path, file_type="docx", max_chunk_chars=500, min_chunk_chars=1)

    ids1 = [c.chunk_id for c in out1.chunks]
    ids2 = [c.chunk_id for c in out2.chunks]
    assert ids1 == ids2


def test_empty_docx_raises(tmp_path: Path) -> None:
    path = tmp_path / "empty.docx"
    _write_docx(path, with_content=False)

    with pytest.raises(EmptyDocumentError):
        parse_and_chunk(path, file_type="docx")


def test_unsupported_type_raises(tmp_path: Path) -> None:
    path = tmp_path / "x.pdf"
    path.write_text("not really a pdf", encoding="utf-8")

    with pytest.raises(ValueError):
        parse_and_chunk(path)
