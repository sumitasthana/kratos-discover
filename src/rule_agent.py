from __future__ import annotations

import json
import logging
import re
from datetime import datetime
from dataclasses import dataclass
from enum import Enum
from pathlib import Path
from typing import Any, Dict, List, Optional, TypedDict

from pydantic import BaseModel, Field, ValidationError
from pydantic.config import ConfigDict

from langgraph.graph import END, StateGraph

from prompt_registry import PromptRegistry

logger = logging.getLogger(__name__)


class RuleCategory(str, Enum):
    RULE = "rule"
    CONTROL = "control"
    RISK = "risk"


class RuleType(str, Enum):
    DATA_QUALITY_THRESHOLD = "data_quality_threshold"
    OWNERSHIP_CATEGORY = "ownership_category"
    BENEFICIAL_OWNERSHIP_THRESHOLD = "beneficial_ownership_threshold"
    DOCUMENTATION_REQUIREMENT = "documentation_requirement"
    UPDATE_REQUIREMENT = "update_requirement"
    UPDATE_TIMELINE = "update_timeline"
    CONTROL_REQUIREMENT = "control_requirement"
    RISK_STATEMENT = "risk_statement"


class RuleMetadata(BaseModel):
    source_block: str
    source_location: str


class Rule(BaseModel):
    rule_id: str
    category: RuleCategory
    rule_type: RuleType
    rule_description: str
    grounded_in: str
    confidence: float = Field(ge=0.5, le=0.99)
    attributes: Dict[str, Any]
    metadata: RuleMetadata


class RulesResponse(BaseModel):
    rules: List[Rule]


class GRCComponentMetadata(BaseModel):
    model_config = ConfigDict(extra="allow")
    source_block: str
    source_location: str


class PolicyComponent(BaseModel):
    model_config = ConfigDict(extra="allow")
    component_type: str = Field(default="policy")
    component_id: Optional[str] = None
    component_title: Optional[str] = None
    component_owner: Optional[str] = None
    policy_objective: Optional[str] = None
    approval_authority: Optional[str] = None
    effective_date: Optional[str] = None
    review_cycle: Optional[str] = None
    policy_statement: Optional[str] = None
    scope: Optional[str] = None
    detailed_requirements: Optional[str] = None
    roles_responsibilities: Optional[str] = None
    related_regulations: Any = None
    grc_platform_module: Optional[str] = None
    related_controls: Any = None
    related_risks: Any = None
    source_table_identifier: Optional[str] = None
    validation_errors: List[str] = Field(default_factory=list)
    metadata: GRCComponentMetadata


class RiskComponent(BaseModel):
    model_config = ConfigDict(extra="allow")
    component_type: str = Field(default="risk")
    component_id: Optional[str] = None
    risk_description: Optional[str] = None
    risk_owner: Optional[str] = None
    risk_category: Optional[str] = None
    inherent_risk_rating: Optional[str] = None
    residual_risk_rating: Optional[str] = None
    effective_date: Optional[str] = None
    review_cycle: Optional[str] = None
    grc_platform_module: Optional[str] = None
    related_policies: Any = None
    mitigation_controls: Any = None
    related_controls: Any = None
    source_table_identifier: Optional[str] = None
    validation_errors: List[str] = Field(default_factory=list)
    metadata: GRCComponentMetadata


class ControlComponent(BaseModel):
    model_config = ConfigDict(extra="allow")
    component_type: str = Field(default="control")
    component_id: Optional[str] = None
    control_description: Optional[str] = None
    control_owner: Optional[str] = None
    control_type: Any = None
    operating_frequency: Optional[str] = None
    testing_frequency: Optional[str] = None
    evidence: Any = None
    effective_date: Optional[str] = None
    review_cycle: Optional[str] = None
    grc_platform_module: Optional[str] = None
    related_policies: Any = None
    related_risks: Any = None
    source_table_identifier: Optional[str] = None
    validation_errors: List[str] = Field(default_factory=list)
    metadata: GRCComponentMetadata


class GRCComponentsResponse(BaseModel):
    policies: List[PolicyComponent] = Field(default_factory=list)
    risks: List[RiskComponent] = Field(default_factory=list)
    controls: List[ControlComponent] = Field(default_factory=list)
    extraction_summary: Dict[str, Any] = Field(default_factory=dict)


@dataclass(frozen=True)
class DocumentSection:
    header: str
    content: str
    source_location: str


class AgentState(TypedDict, total=False):
    document_text: str
    document_path: str
    sections: List[DocumentSection]
    raw_rules: List[Dict[str, Any]]
    validated_rules: List[Dict[str, Any]]
    deduplicated_rules: List[Dict[str, Any]]
    final_rules: List[Rule]

    raw_components: Dict[str, Any]
    validated_components: Dict[str, Any]
    final_components: Dict[str, Any]


class RuleAgent:
    def __init__(
        self,
        registry: Optional[PromptRegistry] = None,
        llm: Any = None,
        fdic_source_path: Optional[str] = None,
    ) -> None:
        self.registry = registry or PromptRegistry(base_dir=Path(__file__).resolve().parent)
        self.llm = llm
        self.fdic_source_path = fdic_source_path
        self.graph = self._build_graph()

    def _build_graph(self):
        graph = StateGraph(AgentState)

        graph.add_node("segment", self._segment_requirements_node)
        graph.add_node("extract", self._extract_rules_node)
        graph.add_node("validate", self._validate_parse_node)
        graph.add_node("deduplicate", self._deduplication_node)
        graph.add_node("ground", self._grounding_scoring_node)

        graph.set_entry_point("segment")
        graph.add_edge("segment", "extract")
        graph.add_edge("extract", "validate")
        graph.add_edge("validate", "deduplicate")
        graph.add_edge("deduplicate", "ground")
        graph.add_edge("ground", END)

        return graph.compile()

    def extract_rules(
        self,
        document_text: Optional[str] = None,
        document_path: Optional[str] = None,
    ) -> List[Rule]:
        if not document_text and not document_path:
            raise ValueError("Provide document_text or document_path")

        state: AgentState = {}
        if document_text:
            state["document_text"] = document_text
        if document_path:
            state["document_path"] = document_path

        result = self.graph.invoke(state)
        return list(result.get("final_rules", []))

    def extract_grc_components(
        self,
        document_text: Optional[str] = None,
        document_path: Optional[str] = None,
    ) -> Dict[str, Any]:
        if not document_text and not document_path:
            raise ValueError("Provide document_text or document_path")

        state: AgentState = {}
        if document_text:
            state["document_text"] = document_text
        if document_path:
            state["document_path"] = document_path

        state = self._segment_requirements_node(state)
        state = self._extract_grc_components_node(state)
        state = self._validate_grc_components_node(state)
        state = self._ground_grc_components_node(state)
        return dict(state.get("final_components", {}))

    def _segment_requirements_node(self, state: AgentState) -> AgentState:
        if "document_path" in state and state["document_path"]:
            path = Path(state["document_path"])
            text, sections = self._load_and_segment_document(path)
            state["document_text"] = text
            state["sections"] = sections
            return state

        text = state.get("document_text", "")
        if not text:
            state["sections"] = []
            return state

        sections: List[DocumentSection] = []
        lines = text.splitlines()
        buf: List[str] = []
        header = "Document"
        start_line = 1

        def flush(end_line: int) -> None:
            nonlocal buf, header, start_line
            content = "\n".join(buf).strip()
            if content:
                sections.append(
                    DocumentSection(
                        header=header,
                        content=content,
                        source_location=f"lines:{start_line}-{end_line}",
                    )
                )
            buf = []

        for i, line in enumerate(lines, start=1):
            if re.match(r"^\s*#{1,4}\s+", line):
                flush(i - 1)
                header = line.strip().lstrip("#").strip()
                start_line = i
                continue

            buf.append(line)

        flush(len(lines))

        state["sections"] = sections
        return state

    def _extract_rules_node(self, state: AgentState) -> AgentState:
        if self.llm is None:
            raise RuntimeError("LLM is not configured. Pass llm=... to RuleAgent.")

        sections = state.get("sections", [])
        prompt_version = self.registry.get_active_prompt("rule_extraction")
        system_prompt = prompt_version["content"]
        spec = prompt_version.get("spec", {})
        user_template = str(spec.get("user_message_template", "{block_header}\n{block_content}"))

        raw_rules: List[Dict[str, Any]] = []

        for section in sections:
            if ":table" in section.source_location:
                continue
            user_message = self._render_user_message(
                template=user_template,
                block_header=section.header,
                block_content=section.content,
            )

            rules = self._call_llm_rules(system_prompt=system_prompt, user_message=user_message)

            if not isinstance(rules, list):
                logger.warning("LLM returned non-list rules; skipping section %s", section.source_location)
                continue

            for r in rules:
                if not isinstance(r, dict):
                    continue
                r.setdefault("metadata", {})
                # Authoritative source location for downstream grounding.
                r["metadata"]["source_location"] = section.source_location
                # Ensure strict grounding uses the authoritative section content rather than
                # any paraphrase the model might return.
                r["metadata"]["source_block"] = section.content
                raw_rules.append(r)

        state["raw_rules"] = raw_rules
        return state

    def _call_llm_rules(self, system_prompt: str, user_message: str) -> List[Dict[str, Any]]:
        structured = self._call_llm_structured(system_prompt=system_prompt, user_message=user_message)
        if structured is not None:
            return [r.model_dump() for r in structured.rules]

        response_text = self._call_llm_json(system_prompt=system_prompt, user_message=user_message)
        payload = self._safe_json_load(response_text)
        rules = payload.get("rules", []) if isinstance(payload, dict) else []
        return rules if isinstance(rules, list) else []

    def _call_llm_structured(self, system_prompt: str, user_message: str) -> Optional[RulesResponse]:
        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_message},
        ]

        try:
            if not hasattr(self.llm, "with_structured_output"):
                return None

            try:
                structured_llm = self.llm.with_structured_output(RulesResponse, method="function_calling")
            except TypeError:
                structured_llm = self.llm.with_structured_output(RulesResponse)
            resp = structured_llm.invoke(messages)

            if isinstance(resp, RulesResponse):
                return resp
            if isinstance(resp, dict):
                return RulesResponse.model_validate(resp)
            return None
        except Exception:
            logger.info("Structured output call failed; falling back to JSON parsing", exc_info=True)
            return None

    def _call_llm_grc_components(self, system_prompt: str, user_message: str) -> Dict[str, Any]:
        structured = self._call_llm_grc_structured(system_prompt=system_prompt, user_message=user_message)
        if structured is not None:
            return structured.model_dump()

        response_text = self._call_llm_json(system_prompt=system_prompt, user_message=user_message)
        payload = self._safe_json_load(response_text)
        return payload if isinstance(payload, dict) else {}

    def _call_llm_grc_structured(
        self, system_prompt: str, user_message: str
    ) -> Optional[GRCComponentsResponse]:
        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_message},
        ]

        try:
            if not hasattr(self.llm, "with_structured_output"):
                return None

            try:
                structured_llm = self.llm.with_structured_output(
                    GRCComponentsResponse, method="function_calling"
                )
            except TypeError:
                structured_llm = self.llm.with_structured_output(GRCComponentsResponse)

            resp = structured_llm.invoke(messages)
            if isinstance(resp, GRCComponentsResponse):
                return resp
            if isinstance(resp, dict):
                return GRCComponentsResponse.model_validate(resp)
            return None
        except Exception:
            logger.debug("Structured output call failed; falling back to JSON parsing", exc_info=True)
            return None

    def _render_user_message(self, template: str, block_header: str, block_content: str) -> str:
        rendered = template
        rendered = rendered.replace("{block_header}", block_header)
        rendered = rendered.replace("{block_content}", block_content)
        return rendered

    def _validate_parse_node(self, state: AgentState) -> AgentState:
        raw_rules = state.get("raw_rules", [])
        validated: List[Dict[str, Any]] = []

        for idx, r in enumerate(raw_rules, start=1):
            try:
                model = Rule.model_validate(r)
                validated.append(model.model_dump())
            except ValidationError as e:
                logger.info("Dropping invalid rule #%s: %s", idx, e)

        state["validated_rules"] = validated
        return state

    def _extract_grc_components_node(self, state: AgentState) -> AgentState:
        if self.llm is None:
            raise RuntimeError("LLM is not configured. Pass llm=... to RuleAgent.")

        sections = state.get("sections", [])
        table_sections = [s for s in sections if ":table" in s.source_location]
        if table_sections:
            sections = table_sections
        prompt_version = self.registry.get_active_prompt("grc_component_extraction")
        system_prompt = prompt_version["content"]
        spec = prompt_version.get("spec", {})
        user_template = str(spec.get("user_message_template", "{block_header}\n{block_content}"))

        aggregated: Dict[str, Any] = {
            "policies": [],
            "risks": [],
            "controls": [],
            "extraction_summary": {},
        }

        for section in sections:
            user_message = self._render_user_message(
                template=user_template,
                block_header=section.header,
                block_content=section.content,
            )

            payload = self._call_llm_grc_components(system_prompt=system_prompt, user_message=user_message)
            if not isinstance(payload, dict):
                continue

            for key in ("policies", "risks", "controls"):
                items = payload.get(key, [])
                if not isinstance(items, list):
                    continue
                for item in items:
                    if not isinstance(item, dict):
                        continue
                    item.setdefault("metadata", {})
                    item["metadata"]["source_location"] = section.source_location
                    item["metadata"]["source_block"] = section.content
                    aggregated[key].append(item)

        state["raw_components"] = aggregated
        return state

    def _validate_grc_components_node(self, state: AgentState) -> AgentState:
        raw = state.get("raw_components", {})
        validated: Dict[str, Any] = {
            "policies": [],
            "risks": [],
            "controls": [],
            "extraction_summary": dict(raw.get("extraction_summary", {})) if isinstance(raw, dict) else {},
        }

        def _append_error(model: BaseModel, msg: str) -> None:
            existing = getattr(model, "validation_errors", [])
            if not isinstance(existing, list):
                existing = []
            errors = list(existing)
            errors.append(msg)
            setattr(model, "validation_errors", errors)

        def _normalize_list(val: Any) -> List[str]:
            if val is None:
                return []
            if isinstance(val, list):
                out: List[str] = []
                for x in val:
                    if x is None:
                        continue
                    s = str(x).strip()
                    if s:
                        out.append(s)
                return out
            if isinstance(val, str):
                s = val.strip()
                if not s:
                    return []
                parts = re.split(r"\s*(?:,|;|\n|\||\u2022|\u25cf)\s*", s)
                return [p.strip() for p in parts if p and p.strip()]
            return [str(val).strip()] if str(val).strip() else []

        def _normalize_list_field(model: BaseModel, field: str) -> None:
            try:
                current = getattr(model, field, None)
            except Exception:
                return
            normalized = _normalize_list(current)
            setattr(model, field, normalized)

        def _try_parse_date_to_iso(date_str: str) -> Optional[str]:
            s = date_str.strip()
            if not s:
                return None
            # Common formats in FDIC/Archer exports.
            formats = [
                "%Y-%m-%d",
                "%m/%d/%Y",
                "%B %d, %Y",
                "%b %d, %Y",
                "%B %d %Y",
                "%b %d %Y",
            ]
            for fmt in formats:
                try:
                    return datetime.strptime(s, fmt).date().isoformat()
                except ValueError:
                    continue
            return None

        def _normalize_date_field(model: BaseModel, field: str) -> None:
            val = getattr(model, field, None)
            if val is None:
                return
            if not isinstance(val, str):
                return
            s = val.strip()
            if not s:
                return
            iso = _try_parse_date_to_iso(s)
            if iso is None:
                # Keep original string per requirement.
                _append_error(model, f"date_unparseable: {field}")
                return
            setattr(model, field, iso)

        def _normalize_control_type_field(model: BaseModel, field: str = "control_type") -> None:
            val = getattr(model, field, None)
            if val is None:
                return
            if isinstance(val, str):
                return
            if isinstance(val, dict):
                nature = str(val.get("nature") or "").strip()
                automation = str(val.get("automation") or "").strip()
                raw = str(val.get("raw") or "").strip()
                if nature and automation:
                    setattr(model, field, f"{nature} / {automation}")
                    return
                if raw:
                    setattr(model, field, raw)
                    return
                setattr(model, field, json.dumps(val, ensure_ascii=False))
                return
            setattr(model, field, str(val))

        def _flag_missing_required(model: BaseModel, required_fields: List[str]) -> None:
            existing = getattr(model, "validation_errors", [])
            if not isinstance(existing, list):
                existing = []

            errors = list(existing)
            for field in required_fields:
                val = getattr(model, field, None)
                if val is None or (isinstance(val, str) and not val.strip()):
                    errors.append(f"missing_required_field: {field}")

            setattr(model, "validation_errors", errors)

        for p in list(raw.get("policies", [])) if isinstance(raw, dict) else []:
            if not isinstance(p, dict):
                continue
            try:
                model = PolicyComponent.model_validate(p)
                _normalize_list_field(model, "related_regulations")
                _normalize_list_field(model, "related_controls")
                _normalize_list_field(model, "related_risks")
                _normalize_date_field(model, "effective_date")
                _flag_missing_required(
                    model,
                    required_fields=["component_id", "component_title", "component_owner", "policy_objective"],
                )
                validated["policies"].append(model)
            except ValidationError as e:
                logger.info("Dropping invalid policy component: %s", e)

        for r in list(raw.get("risks", [])) if isinstance(raw, dict) else []:
            if not isinstance(r, dict):
                continue
            try:
                model = RiskComponent.model_validate(r)
                _normalize_list_field(model, "related_policies")
                _normalize_list_field(model, "mitigation_controls")
                _normalize_list_field(model, "related_controls")
                _normalize_date_field(model, "effective_date")
                _flag_missing_required(model, required_fields=["component_id", "risk_description", "risk_owner"])
                validated["risks"].append(model)
            except ValidationError as e:
                logger.info("Dropping invalid risk component: %s", e)

        for c in list(raw.get("controls", [])) if isinstance(raw, dict) else []:
            if not isinstance(c, dict):
                continue
            try:
                model = ControlComponent.model_validate(c)
                _normalize_list_field(model, "evidence")
                _normalize_list_field(model, "related_policies")
                _normalize_list_field(model, "related_risks")
                _normalize_date_field(model, "effective_date")
                _normalize_control_type_field(model)
                _flag_missing_required(
                    model,
                    required_fields=["component_id", "control_description", "control_owner"],
                )
                validated["controls"].append(model)
            except ValidationError as e:
                logger.info("Dropping invalid control component: %s", e)

        state["validated_components"] = validated
        return state

    def _ground_grc_components_node(self, state: AgentState) -> AgentState:
        # For v1.2 components, we keep strict traceability by attaching authoritative
        # section source_block/source_location at extraction time. We do not
        # aggressively drop components here because table extraction often yields
        # terse identifiers.
        validated = state.get("validated_components", {})
        policies: List[PolicyComponent] = list(validated.get("policies", []))
        risks: List[RiskComponent] = list(validated.get("risks", []))
        controls: List[ControlComponent] = list(validated.get("controls", []))

        policy_ids = {p.component_id for p in policies if isinstance(p, PolicyComponent) and p.component_id}
        risk_ids = {r.component_id for r in risks if isinstance(r, RiskComponent) and r.component_id}
        control_ids = {c.component_id for c in controls if isinstance(c, ControlComponent) and c.component_id}

        def _append_error(model: BaseModel, msg: str) -> None:
            existing = getattr(model, "validation_errors", [])
            if not isinstance(existing, list):
                existing = []
            errors = list(existing)
            errors.append(msg)
            setattr(model, "validation_errors", errors)

        def _check_refs(model: BaseModel, field: str, allowed: set[str]) -> None:
            val = getattr(model, field, None)
            if not val:
                return
            if isinstance(val, str):
                vals = [v.strip() for v in re.split(r"\s*(?:,|;|\n|\|)\s*", val) if v.strip()]
            elif isinstance(val, list):
                vals = [str(v).strip() for v in val if v is not None and str(v).strip()]
            else:
                vals = [str(val).strip()] if str(val).strip() else []
            for v in vals:
                if v and v not in allowed:
                    _append_error(model, f"missing_reference: {field}={v}")

        for p in policies:
            if not isinstance(p, PolicyComponent):
                continue
            _check_refs(p, "related_controls", control_ids)
            _check_refs(p, "related_risks", risk_ids)

        for r in risks:
            if not isinstance(r, RiskComponent):
                continue
            _check_refs(r, "related_policies", policy_ids)
            _check_refs(r, "mitigation_controls", control_ids)
            _check_refs(r, "related_controls", control_ids)

        for c in controls:
            if not isinstance(c, ControlComponent):
                continue
            _check_refs(c, "related_policies", policy_ids)
            _check_refs(c, "related_risks", risk_ids)

        tables_processed = len(
            {
                getattr(m, "metadata").source_location
                for m in [*policies, *risks, *controls]
                if hasattr(m, "metadata") and getattr(m, "metadata") is not None
            }
        )

        all_errors: List[str] = []
        warnings: List[str] = []
        for m in [*policies, *risks, *controls]:
            errs = getattr(m, "validation_errors", [])
            if isinstance(errs, list):
                all_errors.extend([str(e) for e in errs])

        if any(e.startswith("missing_reference") for e in all_errors):
            warnings.append("One or more cross-references could not be validated")

        # Simple completeness score: fraction of required fields present.
        required_map: List[tuple[str, List[str], List[BaseModel]]] = [
            ("policy", ["component_id", "component_title", "component_owner", "policy_objective"], policies),
            ("risk", ["component_id", "risk_description", "risk_owner"], risks),
            ("control", ["component_id", "control_description", "control_owner"], controls),
        ]
        required_total = 0
        required_present = 0
        for _, fields, models in required_map:
            for m in models:
                for f in fields:
                    required_total += 1
                    v = getattr(m, f, None)
                    if v is None:
                        continue
                    if isinstance(v, str) and not v.strip():
                        continue
                    required_present += 1

        extraction_quality = (required_present / required_total) if required_total else 0.0

        summary = dict(validated.get("extraction_summary", {}))
        summary.update(
            {
                "total_components": len(policies) + len(risks) + len(controls),
                "policies_count": len(policies),
                "risks_count": len(risks),
                "controls_count": len(controls),
                "tables_processed": tables_processed,
                "extraction_quality": round(extraction_quality, 4),
                "validation_errors": all_errors,
                "warnings": warnings,
            }
        )

        final_components: Dict[str, Any] = {
            "policies": policies,
            "risks": risks,
            "controls": controls,
            "extraction_summary": summary,
        }
        state["final_components"] = final_components
        return state

    def _deduplication_node(self, state: AgentState) -> AgentState:
        validated = state.get("validated_rules", [])
        seen: set[str] = set()
        deduped: List[Dict[str, Any]] = []

        for r in validated:
            key = self._dedupe_key(r)
            if key in seen:
                continue
            seen.add(key)
            deduped.append(r)

        state["deduplicated_rules"] = deduped
        return state

    def _grounding_scoring_node(self, state: AgentState) -> AgentState:
        deduped = state.get("deduplicated_rules", [])
        sections = state.get("sections", [])

        by_location: Dict[str, DocumentSection] = {s.source_location: s for s in sections}

        final: List[Rule] = []
        for r in deduped:
            try:
                rule = Rule.model_validate(r)
            except ValidationError:
                continue

            loc = rule.metadata.source_location
            section = by_location.get(loc)
            if section is None:
                continue

            if not self._is_grounded(rule=rule, section=section):
                continue

            final.append(rule)

        state["final_rules"] = final
        return state

    def _is_grounded(self, rule: Rule, section: DocumentSection) -> bool:
        grounded_in = rule.grounded_in.strip()
        source_block = rule.metadata.source_block.strip()
        if not grounded_in or not source_block:
            return False

        section_content = section.content
        if grounded_in not in section_content:
            grounded_norm = self._normalize_ws(grounded_in)
            section_norm = self._normalize_ws(section_content)
            if grounded_norm not in section_norm:
                return False

        # source_block should be the full section content (authoritative). Keep a defensive check.
        if source_block and source_block not in section_content:
            source_norm = self._normalize_ws(source_block)
            section_norm = self._normalize_ws(section_content)
            if source_norm not in section_norm:
                return False

        return True

    def _normalize_ws(self, text: str) -> str:
        return re.sub(r"\s+", " ", text).strip()

    def _dedupe_key(self, r: Dict[str, Any]) -> str:
        desc = str(r.get("rule_description", "")).strip().lower()
        desc = re.sub(r"\s+", " ", desc)
        rule_type = str(r.get("rule_type", "")).strip().lower()
        category = str(r.get("category", "")).strip().lower()
        loc = str(((r.get("metadata") or {}).get("source_location", ""))).strip().lower()
        return "|".join([category, rule_type, loc, desc])

    def _call_llm_json(self, system_prompt: str, user_message: str) -> str:
        try:
            messages = [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_message},
            ]

            if hasattr(self.llm, "invoke"):
                resp = self.llm.invoke(messages)
                if isinstance(resp, str):
                    return resp
                content = getattr(resp, "content", None)
                if isinstance(content, str):
                    return content
                if isinstance(content, list) and content and isinstance(content[0], dict) and "text" in content[0]:
                    return str(content[0]["text"])
                return str(resp)

            return str(self.llm(messages))
        except Exception:
            logger.exception("LLM call failed")
            raise

    def _safe_json_load(self, text: str) -> Dict[str, Any]:
        text = text.strip()
        try:
            return json.loads(text)
        except json.JSONDecodeError:
            match = re.search(r"\{.*\}", text, re.DOTALL)
            if not match:
                return {"rules": []}
            try:
                return json.loads(match.group(0))
            except json.JSONDecodeError:
                return {"rules": []}

    def _load_and_segment_document(self, path: Path) -> tuple[str, List[DocumentSection]]:
        suffix = path.suffix.lower()
        if suffix == ".docx":
            return self._load_docx(path)
        if suffix == ".pdf":
            return self._load_pdf(path)
        if suffix in {".html", ".htm"}:
            return self._load_html(path)
        raise ValueError(f"Unsupported document format: {suffix}")

    def _load_docx(self, path: Path) -> tuple[str, List[DocumentSection]]:
        from docx import Document

        doc = Document(str(path))
        lines: List[str] = []
        sections: List[DocumentSection] = []

        current_header = "FDIC 370"
        current_buf: List[str] = []
        para_index = 0

        def flush_paragraphs() -> None:
            nonlocal current_buf
            content = "\n".join(current_buf).strip()
            if content:
                sections.append(
                    DocumentSection(
                        header=current_header,
                        content=content,
                        source_location=f"docx:{path.name}:p{para_index}",
                    )
                )
            current_buf = []

        for p in doc.paragraphs:
            para_index += 1
            text = (p.text or "").strip()
            if not text:
                continue

            style = getattr(p.style, "name", "") if getattr(p, "style", None) else ""
            if style and style.lower().startswith("heading"):
                flush_paragraphs()
                current_header = text
                current_buf.append(text)
                lines.append(text)
                continue

            current_buf.append(text)
            lines.append(text)

            if len(current_buf) >= 20:
                flush_paragraphs()

        flush_paragraphs()

        for table_idx, table in enumerate(doc.tables, start=1):
            rows: List[Dict[str, Any]] = []
            raw_rows: List[List[str]] = []
            for row in table.rows:
                cells = [((c.text or "").strip()) for c in row.cells]
                raw_rows.append(cells)
                if not any(cells):
                    continue

                field = cells[0] if len(cells) >= 1 else ""
                value = " | ".join(c for c in cells[1:] if c) if len(cells) > 1 else ""
                rows.append({"field": field, "value": value, "cells": cells})

            payload: Dict[str, Any] = {
                "content_type": "docx_table",
                "table_index": table_idx,
                "header": current_header,
                "rows": rows,
                "raw_rows": raw_rows,
                "row_count": len(rows),
            }

            sections.append(
                DocumentSection(
                    header=current_header,
                    content=json.dumps(payload, ensure_ascii=False),
                    source_location=f"docx:{path.name}:table{table_idx}",
                )
            )

        return "\n".join(lines), sections

    def _load_pdf(self, path: Path) -> tuple[str, List[DocumentSection]]:
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        sections: List[DocumentSection] = []
        lines: List[str] = []

        for page_idx, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            text = text.strip()
            if not text:
                continue
            lines.append(text)
            sections.append(
                DocumentSection(
                    header=f"PDF page {page_idx}",
                    content=text,
                    source_location=f"pdf:{path.name}:page{page_idx}",
                )
            )

        return "\n\n".join(lines), sections

    def _load_html(self, path: Path) -> tuple[str, List[DocumentSection]]:
        from bs4 import BeautifulSoup

        html = path.read_text(encoding="utf-8", errors="ignore")
        soup = BeautifulSoup(html, "lxml")

        text = soup.get_text("\n")
        text = re.sub(r"\n{3,}", "\n\n", text).strip()

        sections: List[DocumentSection] = [
            DocumentSection(
                header="HTML document",
                content=text,
                source_location=f"html:{path.name}",
            )
        ]

        return text, sections
